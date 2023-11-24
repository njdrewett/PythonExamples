#! Python3

#PDF to Word converter

# Converts PDf documents to Word documents. Using the installed word document application **REQUIRED**

import win32com.client # install with "pip install pywin32
import docx

wordFilename = 'ouputWordDoucment.docx'
pdfFilename = 'inputPdfDocument.pdf'

doc = docx.Document()
doc.add_paragraph('Hello Converter!')
doc.add_heading('Header 1', 1)
paraObj1 = doc.add_paragraph('This is a second paragraph.')

paraObj2 = doc.add_paragraph('This is a yet another paragraph.')

paraObj1.add_run(' This text is being added to the second paragraph.')
doc.save(wordFilename)

wdFormatPDF = 17 # Word's numeric code for PDFs.
wordObj = win32com.client.Dispatch('Word.Application')
docObj = wordObj.Documents.Open(wordFilename)
docObj.SaveAs(pdfFilename, FileFormat=wdFormatPDF)
docObj.Close()
wordObj.Quit()




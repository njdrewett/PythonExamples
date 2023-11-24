#! python3
# Word documents

import docx

doc = docx.Document()

doc.add_paragraph('Hello world!')
doc.add_heading('Header 0', 0)
paraObj1 = doc.add_paragraph('This is a second paragraph.')

paraObj2 = doc.add_paragraph('This is a yet another paragraph.')

paraObj1.add_run(' This text is being added to the second paragraph.')

doc.save('multipleParagraphs.docx')
